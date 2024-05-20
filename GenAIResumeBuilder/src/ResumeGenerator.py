from docx import Document
from docx.shared import Pt, RGBColor
import json
import os
from src.constants import Config


class DocumentGenerator:
    def replace_text(self, docx_file, text_to_be_replaced, replacing_text):
        for paragraph in docx_file.paragraphs:
            if text_to_be_replaced in paragraph.text:
                paragraph.text = paragraph.text.replace(text_to_be_replaced, replacing_text)
        return docx_file

    def add_heading(self, docx_file, heading):
        paragraph = docx_file.add_paragraph()
        run = paragraph.add_run(heading)

        paragraph.paragraph_format.line_spacing = Config.HEADING_LINE_SPACING
        paragraph.paragraph_format.left_indent = Pt(Config.HEADING_INDENT_LEVEL)

        run.font.bold = Config.HEADING_FONT_BOLD
        run.font.name = Config.HEADING_FONT_NAME
        run.font.size = Pt(Config.HEADING_FONT_SIZE)
        run.font.color.rgb = RGBColor(*Config.HEADING_FONT_COLOR)

    def add_paragraph(self, docx_file, text):
        # print("Adding paragraph : " + text)
        paragraph = docx_file.add_paragraph()
        run = paragraph.add_run(text)

        paragraph.paragraph_format.line_spacing = Config.BODY_LINE_SPACING
        paragraph.paragraph_format.left_indent = Pt(Config.BODY_INDENT_LEVEL)

        run.font.name = Config.BODY_FONT_NAME
        run.font.size = Pt(Config.BODY_FONT_SIZE)
        run.font.color.rgb = RGBColor(*Config.BODY_FONT_COLOR)

    def add_text(self, paragraph, text, is_bold=False, is_italic=False, is_underlined=False):
        # print("Adding text : " + text)
        run = paragraph.add_run(text)
        run.bold = is_bold
        run.underline = is_underlined
        run.font.color.rgb = RGBColor(*Config.BODY_FONT_COLOR)
        run.font.name = Config.BODY_FONT_NAME
        run.font.size = Pt(Config.BODY_FONT_SIZE)

    def add_bulleted_list(self, docx_file, list):
        # print("Adding list : ", list)
        # print(type(list))
        for point in list:
            paragraph = docx_file.add_paragraph()
            paragraph.paragraph_format.line_spacing = Config.BODY_LINE_SPACING
            paragraph.paragraph_format.left_indent = Pt(Config.BODY_INDENT_LEVEL)

            run = paragraph.add_run("\u2022 " + str(point))
            run.font.name = Config.BODY_FONT_NAME
            run.font.size = Pt(Config.BODY_FONT_SIZE)
            run.font.color.rgb = RGBColor(*Config.BODY_FONT_COLOR)

    def add_table(self, docx_file, rows, header_row):
        table = docx_file.add_table(rows=1, cols=len(rows[0]))

        table.rows[0].cells = header_row
        for row in rows:
            table.add_row().cells = row

    def save(self, docx_file, path, file_name):
        docx_file.save(path + '/' + file_name)


class ShortResumeGenerator(DocumentGenerator):
    def __init__(self, json_object, summary_type):
        self.json_object = json_object
        self.summary_type = summary_type

        with open(Config.TEMPLATE_PATH, "rb") as file:
            docx_file = Document(file)
        self.docx_file = docx_file

    def add_title(self):
        name = self.json_object.get("NAME", "") or "UNIDENTIFIED"
        designation = self.json_object.get("DESIGNATION", "") or "UNIDENTIFIED"

        self.docx_file = self.replace_text(self.docx_file, "NAME", name.upper())
        self.docx_file = self.replace_text(self.docx_file, "DESIGNATION", designation.upper())

    def add_summary(self):
        summary = self.json_object.get("SUMMARY", []) or []
        summary = summary if isinstance(summary, list) else list(summary.replace("[", "").replace("]", "").split(". "))
        self.add_heading(self.docx_file, "SUMMARY")
        self.add_bulleted_list(self.docx_file, summary)

    def add_work_experience(self):
        work_experiences = self.json_object.get("WORK EXPERIENCE", "") or []
        self.add_heading(self.docx_file, "WORK EXPERIENCE")
        work_experiences = [
            ", ".join(
                [
                    work_experience.get("Company", "") or "",
                    work_experience.get("Duration", "") or "",
                    work_experience.get("Role", "") or ""
                ]
            )
            for work_experience in work_experiences
        ]
        self.add_bulleted_list(self.docx_file, work_experiences)

    def generate(self):
        # Adding NAME and DESIGNATION
        self.add_title()

        # Adding SUMMARY
        self.add_summary()

        # Adding Work Experience
        self.add_work_experience()

        # Adding Skills
        skills = self.json_object.get("SKILLS", "") or ""
        if skills:
            self.add_heading(self.docx_file, "SKILLS")
            self.add_paragraph(self.docx_file, skills)

        # Adding Certifications
        certifications = self.json_object.get("CERTIFICATIONS", []) or []
        if certifications:
            self.add_heading(self.docx_file, "CERTIFICATIONS")
            self.add_bulleted_list(self.docx_file, certifications)

        # Adding Key Achievements
        key_achievements = self.json_object.get("KEY ACHIEVEMENTS", []) or []
        if key_achievements:
            self.add_heading(self.docx_file, "KEY ACHIEVEMENTS")
            self.add_bulleted_list(self.docx_file, key_achievements)

        name = self.json_object.get("NAME", "") or "UNIDENTIFIED"

        self.save(self.docx_file, Config.OUTPUT_FOLDER, f"{name.upper()}_{self.summary_type}.docx")


class LongResumeGenerator(DocumentGenerator):
    def __init__(self, json_object, summary_type):
        self.json_object = json_object
        self.summary_type = summary_type

        with open(Config.TEMPLATE_PATH, "rb") as file:
            docx_file = Document(file)
        self.docx_file = docx_file

    def add_title(self):
        name = self.json_object.get("NAME", "") or "UNIDENTIFIED"
        designation = self.json_object.get("DESIGNATION", "") or "UNIDENTIFIED"

        self.docx_file = self.replace_text(self.docx_file, "NAME", name.upper())
        self.docx_file = self.replace_text(self.docx_file, "DESIGNATION", designation.upper())

    def add_summary(self):
        summary = self.json_object.get("SUMMARY", []) or []
        summary = summary if isinstance(summary, list) else [summary]
        self.add_heading(self.docx_file, "SUMMARY")
        self.add_bulleted_list(self.docx_file, summary)

    def add_work_experience(self):
        work_experiences = self.json_object.get("WORK EXPERIENCE", []) or []

        if work_experiences:
            self.add_heading(self.docx_file, "WORK EXPERIENCE")

            for work_experience in work_experiences:
                # print(work_experience)
                paragraph = self.docx_file.add_paragraph()
                paragraph.paragraph_format.line_spacing = Config.BODY_LINE_SPACING
                paragraph.paragraph_format.line_indent = Pt(Config.BODY_INDENT_LEVEL)
                self.add_text(paragraph, work_experience.get("Duration", "") or "", is_underlined=True, is_bold=True)
                self.add_text(paragraph,
                              f' - {work_experience.get("Role", "") or ""}, {work_experience.get("Company", "") or ""}')

                for key, value in work_experience.items():
                    if key not in ["Duration", "Role", "Company"] and value:
                        value = value if isinstance(value, list) else [value]
                        self.add_text(self.docx_file.add_paragraph(), key, is_bold=True)
                        self.add_bulleted_list(self.docx_file, value)

    def add_education_details(self):
        education_details = self.json_object.get("EDUCATION", [])
        if education_details:
            self.add_heading(self.docx_file, "EDUCATION")
            # self.add_table(self.docx_file, [education_detail.values() for education_detail in education_details], header_row=education_details[0].keys())
            for key, value in education_details[0].items():
                paragraph = self.docx_file.add_paragraph()
                self.add_text(paragraph, key.upper(), is_bold=True)
                self.add_text(paragraph, f' - {value or ""}')

    def generate(self):
        # Adding NAME and DESIGNATION
        self.add_title()

        # Adding SUMMARY
        self.add_summary()

        # Adding Skills
        skills = self.json_object.get("SKILLS", "") or ""
        self.add_heading(self.docx_file, "SKILLS")
        self.add_paragraph(self.docx_file, skills)

        # Adding Work Experience
        self.add_work_experience()

        # Adding Education Details
        self.add_education_details()

        # Adding Awards
        awards = self.json_object.get("AWARDS", []) or []
        if awards:
            self.add_heading(self.docx_file, "AWARDS")
            self.add_bulleted_list(self.docx_file, awards)

        # Adding Certifications
        certifications = self.json_object.get("CERTIFICATIONS", []) or []
        if certifications:
            self.add_heading(self.docx_file, "CERTIFICATIONS")
            self.add_bulleted_list(self.docx_file, certifications)

        # Adding Key Achievements
        key_achievements = self.json_object.get("KEY ACHIEVEMENTS", []) or []
        if key_achievements:
            self.add_heading(self.docx_file, "KEY ACHIEVEMENTS")
            self.add_bulleted_list(self.docx_file, key_achievements)

        name = self.json_object.get("NAME", "") or "UNIDENTIFIED"

        self.save(self.docx_file, Config.OUTPUT_FOLDER, f"{name.upper()}_{self.summary_type}.docx")